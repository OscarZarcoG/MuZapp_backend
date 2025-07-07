from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import os
from django.conf import settings
from django.http import HttpResponse
from datetime import datetime

class ContratoGenerator:
    """
    Generador de contratos en formato DOCX usando python-docx
    Permite crear plantillas de Word con combinación de correspondencia
    """
    
    def __init__(self, contrato):
        self.contrato = contrato
        self.doc = Document()
        self.setup_document()
    
    def setup_document(self):
        """Configura el documento con márgenes y estilos"""
        # Configurar márgenes
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
    
    def add_header_with_logo(self):
        """Añade el encabezado con logo y información del grupo"""
        # Crear tabla para el encabezado
        header_table = self.doc.add_table(rows=1, cols=2)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Celda izquierda para el logo
        logo_cell = header_table.cell(0, 0)
        logo_paragraph = logo_cell.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Aquí se puede agregar el logo si existe
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture('media/logo/Logotipo.jpg', width=Inches(1.5))
        
        # Por ahora agregamos texto del logo
        logo_run = logo_paragraph.add_run("GRUPO MUSICAL\nBAHIA MIX")
        logo_run.font.size = Pt(14)
        logo_run.font.bold = True
        
        # Celda derecha para información del contrato
        info_cell = header_table.cell(0, 1)
        info_paragraph = info_cell.paragraphs[0]
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info_run = info_paragraph.add_run("CONTRATO MUSICAL\nGRUPO MUSICAL\n\nBAHIA MIX")
        info_run.font.size = Pt(12)
        info_run.font.bold = True
        
        # Eliminar bordes de la tabla
        self._remove_table_borders(header_table)
        
        # Agregar espacio después del encabezado
        self.doc.add_paragraph()
    
    def add_contract_info(self):
        """Añade la información básica del contrato"""
        # Fecha y lugar
        fecha_paragraph = self.doc.add_paragraph()
        fecha_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fecha_run = fecha_paragraph.add_run(f"Piura de del {datetime.now().year}")
        fecha_run.font.size = Pt(11)
        
        # Información del representante
        self.doc.add_paragraph()
        representante_p = self.doc.add_paragraph()
        representante_run = representante_p.add_run(
            "El Señor Oscar Zarco Solano, identificado con DNI 41648976 Representante del grupo KEE ROLLO y el contratante Señor(a) "
        )
        representante_run.font.size = Pt(11)
        
        # Línea para el nombre del cliente
        cliente_nombre = self.contrato.cliente.nombre_cliente if self.contrato.cliente else "_" * 30
        cliente_run = representante_p.add_run(f"{cliente_nombre}")
        cliente_run.font.size = Pt(11)
        cliente_run.font.underline = True
        
        # Continuar con el texto
        texto_continuo = representante_p.add_run(
            ", Interesado en contratar al grupo, tiene a bien realizar un contrato musical, el contratante solicita los servicios musicales por 5 horas y de conformidad se aceptan las siguientes clausulas"
        )
        texto_continuo.font.size = Pt(11)
    
    def add_contract_clauses(self):
        """Añade las cláusulas del contrato"""
        clauses = [
            {
                "numero": "1.-",
                "texto": f"El contratante el (la) sr (a) {self.contrato.cliente.nombre_cliente if self.contrato.cliente else '_' * 20} al grupo BAHIA MIX, para el día {self.contrato.fecha_evento.strftime('%d/%m/%Y')} a partir de las {self.contrato.hora_inicio.strftime('%H:%M')} hrs y del día {self.contrato.fecha_evento.strftime('%d/%m/%Y')} del mes de {self.contrato.fecha_evento.strftime('%B')} queda comprometido a pagar {self.contrato.pago_total} por las horas pactadas."
            },
            {
                "numero": "2.-",
                "texto": f"El contratante anticipa la cantidad de {self.contrato.pago_adelanto} quedando a cuenta {self.contrato.pago_restante} cantidad que se hará efectiva antes de iniciar."
            },
            {
                "numero": "3.-",
                "texto": "El contrato no podrá transferir de fecha, en caso de suspensión del servicio, por cualquier causa contratante deberá pagar el 50 % de la cantidad acordada"
            },
            {
                "numero": "4.-",
                "texto": f"El domicilio donde se efectuara el evento será en: {self.contrato.nombre_lugar}"
            },
            {
                "numero": "5.-",
                "texto": "El contratante se compromete a proporcionar un lugar suficiente amplio para la ejecución de trabajo del grupo, cubierto para que en caso de lluvia quede cubierto el equipo , además de proporcionar el suministro de energía eléctrica necesaria y suficiente para llevar a cabo el trabajo"
            },
            {
                "numero": "6.-",
                "texto": "Si por alguna situación de riña, pelea o otra situación, se suspende el evento el contratante se compromete al pago del 100% de lo pactado."
            },
            {
                "numero": "7.-",
                "texto": "En caso de alternar con otro grupo musical, mariachi, sonido o alguna variedad, BAHIA MIX no responderá dicho tiempo y se ejecutara el horario de inicio y final del presente contrato."
            },
            {
                "numero": "8.-",
                "texto": "El grupo BAHIA MIX se compromete a realizar su trabajo en el horario y fecha pactados en este contrato y en los términos expresados anteriormente."
            }
        ]
        
        for clause in clauses:
            clause_p = self.doc.add_paragraph()
            clause_p.paragraph_format.space_after = Pt(6)
            
            # Número de la cláusula
            numero_run = clause_p.add_run(clause["numero"])
            numero_run.font.bold = True
            numero_run.font.size = Pt(11)
            
            # Texto de la cláusula
            texto_run = clause_p.add_run(f" {clause['texto']}")
            texto_run.font.size = Pt(11)
        
        # Texto final
        self.doc.add_paragraph()
        final_p = self.doc.add_paragraph("Leídas las clausulas anteriores se firma de conformidad")
        final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        final_run = final_p.runs[0]
        final_run.font.size = Pt(11)
        final_run.font.bold = True
    
    def add_signatures(self):
        """Añade la sección de firmas"""
        # Agregar espacio
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Crear tabla para las firmas
        signature_table = self.doc.add_table(rows=4, cols=2)
        signature_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Líneas de firma
        signature_table.cell(0, 0).text = "_" * 40
        signature_table.cell(0, 1).text = "_" * 40
        
        # Nombres
        signature_table.cell(1, 0).text = "CRISTHIAN FRANK ICARNAQUE LUJAN."
        signature_table.cell(1, 1).text = "CONTRATANTE"
        
        # Títulos
        signature_table.cell(2, 0).text = "Representante del grupo bahia mix"
        signature_table.cell(2, 1).text = "Nombre:"
        
        # DNI
        signature_table.cell(3, 1).text = "DNI:"
        
        # Centrar texto en todas las celdas
        for row in signature_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
        
        # Eliminar bordes
        self._remove_table_borders(signature_table)
    
    def add_contact_info(self):
        """Añade información de contacto"""
        self.doc.add_paragraph()
        
        # Crear tabla para redes sociales
        contact_table = self.doc.add_table(rows=2, cols=2)
        
        # Facebook
        fb_cell = contact_table.cell(0, 0)
        fb_paragraph = fb_cell.paragraphs[0]
        fb_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fb_run = fb_paragraph.add_run("📘 Grupo musical Bahia Mix")
        fb_run.font.size = Pt(10)
        
        # WhatsApp
        wa_cell = contact_table.cell(1, 0)
        wa_paragraph = wa_cell.paragraphs[0]
        wa_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        wa_run = wa_paragraph.add_run("📱 928 267 419")
        wa_run.font.size = Pt(10)
        
        # Eliminar bordes
        self._remove_table_borders(contact_table)
    
    def _remove_table_borders(self, table):
        """Elimina los bordes de una tabla"""
        # Método simplificado para eliminar bordes
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            tblBorders = tblPr.first_child_found_in(qn('w:tblBorders'))
            if tblBorders is not None:
                tblPr.remove(tblBorders)
        except:
            # Si hay algún error, simplemente continuar sin eliminar bordes
            pass
    
    def generate_contract(self):
        """Genera el contrato completo"""
        self.add_header_with_logo()
        self.add_contract_info()
        self.add_contract_clauses()
        self.add_signatures()
        self.add_contact_info()
        
        return self.doc
    
    def save_to_response(self, filename=None):
        """Guarda el documento y lo retorna como HttpResponse"""
        if not filename:
            filename = f"Contrato_{self.contrato.numero_contrato}.docx"
        
        # Generar el documento
        doc = self.generate_contract()
        
        # Crear respuesta HTTP
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        # Guardar documento en memoria
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        response.write(doc_io.getvalue())
        return response
    
    def save_to_file(self, filepath):
        """Guarda el documento en un archivo"""
        doc = self.generate_contract()
        doc.save(filepath)
        return filepath


def generar_contrato_docx(contrato):
    """
    Función helper para generar un contrato en formato DOCX
    """
    generator = ContratoGenerator(contrato)
    return generator.save_to_response()


def generar_contrato_pdf_desde_docx(contrato):
    """
    Genera un PDF a partir del DOCX usando python-docx2pdf
    Requiere tener Microsoft Word instalado o LibreOffice
    """
    try:
        from docx2pdf import convert
        import tempfile
        import pythoncom
        
        # Inicializar COM para Windows
        pythoncom.CoInitialize()
        
        try:
            # Crear archivo temporal para DOCX
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                generator = ContratoGenerator(contrato)
                doc = generator.generate_contract()
                doc.save(temp_docx.name)
                
                # Crear archivo temporal para PDF
                temp_pdf_path = temp_docx.name.replace('.docx', '.pdf')
                
                # Convertir DOCX a PDF
                convert(temp_docx.name, temp_pdf_path)
                
                # Leer el PDF y crear respuesta
                with open(temp_pdf_path, 'rb') as pdf_file:
                    response = HttpResponse(
                        pdf_file.read(),
                        content_type='application/pdf'
                    )
                    response['Content-Disposition'] = f'attachment; filename="Contrato_{contrato.numero_contrato}.pdf"'
                
                # Limpiar archivos temporales
                os.unlink(temp_docx.name)
                os.unlink(temp_pdf_path)
                
                return response
        
        finally:
            # Limpiar COM
            pythoncom.CoUninitialize()
            
    except ImportError:
        # Si no está disponible docx2pdf, retornar el DOCX
        return generar_contrato_docx(contrato)
    except Exception as e:
        # En caso de cualquier otro error, retornar el DOCX
        return generar_contrato_docx(contrato)